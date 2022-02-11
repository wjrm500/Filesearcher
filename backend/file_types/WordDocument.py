from docx import Document

from backend.file_types.File import File

class WordDocument(File):
    extension = 'docx'

    def get_text(self) -> str:
        document = Document(self.filename)
        return '\n'.join([p.text for p in document.paragraphs])